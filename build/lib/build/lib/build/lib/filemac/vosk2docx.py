import json
import pyaudio
from vosk import Model, KaldiRecognizer
import os
from docx import Document

# Load the Vosk model
model_path = "../src/vosk-model-en-us-0.22-lgraph"  # Replace with the path to your Vosk model
model = Model(model_path)
recognizer = KaldiRecognizer(model, 16000)

# Set up the audio stream
audio = pyaudio.PyAudio()
stream = audio.open(format=pyaudio.paInt16, channels=1, rate=16000, input=True, frames_per_buffer=8000)
stream.start_stream()

os.system('clear')
print("\033[1;30mSpeak into the microphone...\033[0m")

# Create a new Word document
doc = Document()
doc.add_heading("Speech Recognition Transcription", level=1)

text = ""

def add_wrapped_text_to_doc(doc, text, width):
    """Add wrapped text to the Word document."""
    wrapped_text = wrap_text(text, width)
    doc.add_paragraph(wrapped_text)

def wrap_text(text, width):
    """Wrap the text to fit within the specified width."""
    wrapped_lines = []
    words = text.split()
    current_line = ""

    for word in words:
        # Check if adding the next word exceeds the width
        if len(current_line) + len(word) + 1 > width:
            wrapped_lines.append(current_line)
            current_line = word  # Start a new line with the current word
        else:
            if current_line:
                current_line += " "  # Add a space before the next word
            current_line += word

    if current_line:
        wrapped_lines.append(current_line)  # Add any remaining text

    return "\n".join(wrapped_lines)

try:
    while True:
        data = stream.read(4096)
        if recognizer.AcceptWaveform(data):
            result = json.loads(recognizer.Result())
            text += " " + result['text']

            # Clear and add the latest wrapped text to the document
            doc.add_paragraph(result['text'])

        else:
            partial_result = json.loads(recognizer.PartialResult())
            partial_text = partial_result['partial']

except KeyboardInterrupt:
    print("\nExiting...")
finally:
    # Save the document
    doc.save("transcription.docx")
    stream.stop_stream()
    stream.close()
    audio.terminate()
