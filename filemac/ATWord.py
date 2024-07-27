from docx import Document
from docx.shared import Pt, RGBColor
# import os


class AdvancedT2word:
    def __init__(self, obj, out_obj=None, fsize=None, fstyle='Times New Roman'):
        self.obj = obj
        self.out_obj = out_obj
        self.fsize = fsize
        self.fstyle = fstyle
        if self.out_obj is None:
            self.out_obj = f"{self.obj.split('.')[0]}_filemac.docx"

    def text_to_word(self):
        # Create a new Document
        doc = Document()

        # Define formatting for headings and body text
        head_font_name = self.fstyle
        heading_styles = {
            # Heading 1
            1: {'font_size': Pt(16), 'font_color': RGBColor(126, 153, 184)},
            # Heading 2
            2: {'font_size': Pt(14), 'font_color': RGBColor(0, 0, 200)},
            # Heading 3
            3: {'font_size': Pt(13), 'font_color': RGBColor(0, 0, 150)},
        }

        body_font_name = 'Times New Roman'
        body_font_size = Pt(12)
        body_font_color = RGBColor(0, 0, 0)  # Black color

        # Open the text file and read content
        with open(self.obj, 'r') as file:
            lines = file.readlines()

        for line in lines:
            # Determine heading level or body text
            if line.startswith('#'):
                level = line.count('#')
                level = min(level, 3)  # Support up to 3 levels of headings
                style = heading_styles.get(level, heading_styles[1])

                p = doc.add_paragraph()
                # Remove '#' and extra space
                run = p.add_run(line[level+1:].strip())
                run.font.size = style['font_size']
                run.font.name = head_font_name
                run.font.color.rgb = style['font_color']
                p.style = f'Heading{level}'
            else:
                p = doc.add_paragraph()
                run = p.add_run(line.strip())
                run.font.name = body_font_name
                run.font.size = body_font_size
                run.font.color.rgb = body_font_color

        # Save the document
        doc.save(self.out_obj)
        print(f"Text file converted to Word document: {self.out_obj}")


if __name__ == "__main__":
    init = AdvancedT2word('/home/skye/Documents/FMAC/file2.txt')

    # Call the function
    init.text_to_word()
