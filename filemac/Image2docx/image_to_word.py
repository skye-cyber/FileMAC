import os
import sys
from PIL import Image
from typing import List, Tuple, Union
from docx import Document
from docx.shared import Inches, Mm
from pathlib import Path
from utils.formats import SUPPORTED_IMAGE_FORMATS
from utils.colors import foreground

fcl = foreground()
RESET = fcl.RESET


class ImageToDocxConverter:
    """
    A class for converting images to DOCX documents.
    """

    def __init__(
        self,
        image_list: list = None,
        input_dir: Union[str, os.PathLike] = None,
        output_path: Union[str, os.PathLike] = None,
        image_size: Tuple[float, float] = (6, 8),  # Default to 6x8 inches
        margin_mm: float = 25,  # Default margin of 25mm (approx 1 inch)
    ) -> None:
        """
        Initializes the ImageToDocxConverter object.

        Args:
            output_path: Path to save the output DOCX file + the file name e.g ~/Document/output.docx.
            filename: Name of the output DOCX file.
            image_size: Tuple (width, height) in inches.
            margin_mm: Margin in millimeters.
        """
        self.image_list = image_list
        self.input_dir = input_dir
        self.output_path = output_path if output_path else self.ensure_output_file()
        self.image_size = image_size
        self.margin_mm = margin_mm
        self.document = Document()  # Create a new document object filename

        # Set document margins in the constructor
        sections = self.document.sections
        for section in sections:
            section.top_margin = Mm(self.margin_mm)
            section.bottom_margin = Mm(self.margin_mm)
            section.left_margin = Mm(self.margin_mm)
            section.right_margin = Mm(self.margin_mm)
        self.create_output_directory()  # Create output directory in constructor

    def ensure_output_file(self) -> os.PathLike:
        file_name = "filemac_image2docx.docx"
        if self.input_dir:
            base_dir = self.input_dir
        else:
            base_dir = Path(self.image_list[0]).parent

        file_path = os.path.join(base_dir, file_name)

        return file_path

    def create_output_directory(self) -> None:
        """
        Creates the output directory if it does not exist.
        """
        Path(self.output_path).parent.mkdir(parents=True, exist_ok=True)

    def get_valid_images(self, image_paths: List[str]) -> List[str]:
        """
        Filters the list of image paths, returning only those with supported formats.

        Args:
            image_paths: A list of file paths to images.

        Returns:
            A list of file paths to valid images.
        """
        valid_images = []
        for image_path in image_paths:
            try:
                if Image.open(image_path).format.lower() in [
                    _formats[1:] for _formats in SUPPORTED_IMAGE_FORMATS.values()
                ]:
                    valid_images.append(image_path)
                else:
                    print(
                        f"{fcl.MAGENTA_FG}Skipping unsupported image format: {fcl.CYAN_FG}{image_path}{RESET}"
                    )
            except Exception as e:
                print(
                    f"{fcl.RED_FG}Error processing image {fcl.YELLOW_FG}{image_path} - {fcl.RED_FG} {e}{RESET}"
                )
        return valid_images

    def convert_images_to_docx(self, image_paths: List[str]) -> os.PathLike:
        """
        Converts a list of images to a single DOCX document.

        Args:
            image_paths: List of image file paths.
        """

        valid_images = self.get_valid_images(image_paths)
        if not valid_images:
            print("No valid images to convert.")
            return

        for image_path in valid_images:
            try:
                # Add a paragraph for each image
                paragraph = self.document.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(
                    image_path,
                    width=Inches(self.image_size[0]),
                    height=Inches(self.image_size[1]),
                )
                # Add a page break after each image, except the last one
                if image_path != valid_images[-1]:
                    self.document.add_page_break()
            except Exception as e:
                print(
                    f"{fcl.RED_FG}Error processing image {fcl.YELLOW_FG}{image_path}:{fcl.RED_FG} {e}{RESET}"
                )

        docx_file_path = (
            self.output_path
            if self.output_path.endswith(("docx", "doc"))
            else f"{self.output_path}.docx"
        )
        self.document.save(docx_file_path)
        return docx_file_path

    def convert_images_in_directory(self, input_dir, output_path) -> os.PathLike:
        """
        Converts all images in a directory to a PDF.

        Args:
            input_dir (str): The directory containing the images.
            output_path (str): The path to save the generated Word File.
            file_extensions (tuple, optional): Tuple of image file extensions to include.
        """

        if not os.path.exists(input_dir):
            raise FileNotFoundError(f"Directory not found: {input_dir}")

        image_paths = sorted(
            [os.path.join(input_dir, f) for f in os.listdir(input_dir)]
        )

        image_paths = self.get_valid_images(image_paths)

        if not image_paths:
            raise ValueError(f"No images found in directory: {input_dir}")

        self.create_pdf_from_images(image_paths, output_path)
        return output_path

    def run(self) -> os.PathLike:
        """
        Runs the conversion process.

        Args:
            image_paths: List of image file paths to convert.
        """
        if not any((self.image_list, self.input_dir)):
            print("No image paths provided.")
            sys.exit()

        if self.image_list and self.output_path:
            if all(os.path.exists(img) for img in self.image_list):
                docx_file_path = self.convert_images_to_docx(self.image_list)
        elif self.input_dir and self.output_path:
            if os.path.exists(self.input_dir):
                docx_file_path = self.convert_images_in_directory(
                    self.input_dir, self.output_path
                )

        if docx_file_path:
            print(
                f"{fcl.GREEN_RG}Successfully created DOCX: {fcl.BLUE_FG}{docx_file_path}{RESET}"
            )

        return docx_file_path


def main(args: List[str]) -> None:
    """
    Main function to parse command line arguments and perform the conversion.

    Args:
        args: List of command line arguments.
    """
    if not args or "-h" in args or "--help" in args:
        print(
            """
            Usage: python image_to_docx.py [options] image1 image2 ... imageN

            Options:
                -h, --help            show this help message and exit
                -o, --output PATH     path to save the output DOCX file (default: current directory)
                -n, --name FILENAME   name of the output DOCX file (default: output_document)
                -s, --size WIDTHxHEIGHT  size of images in inches (e.g., 6x8) (default: 6x8)
                -m, --margin MARGIN_MM margin in millimeters (default: 25)
            """
        )
        sys.exit()

    image_paths = []
    output_path = "."  # Current directory
    filename = "output_document"
    image_size = (6, 8)  # Default 6x8 inches
    margin_mm = 25

    i = 1
    while i < len(args):
        if args[i] in ("-o", "--output"):
            output_path = args[i + 1]
            i += 2
        elif args[i] in ("-n", "--name"):
            filename = args[i + 1]
            i += 2
        elif args[i] in ("-s", "--size"):
            try:
                size_str = args[i + 1]
                width, height = map(float, size_str.split("x"))
                image_size = (width, height)
            except ValueError:
                print("Invalid size format. Please use WIDTHxHEIGHT (e.g., 6x8).")
                sys.exit(1)
            i += 2
        elif args[i] in ("-m", "--margin"):
            try:
                margin_mm = float(args[i + 1])
            except ValueError:
                print("Invalid margin format. Please provide a numeric value.")
                sys.exit(1)
            i += 2
        else:
            if not args[i].startswith("-"):
                image_paths.append(args[i])
                i += 1
            else:
                print(f"Unknown argument: {args[i]}")
                sys.exit(1)

    converter = ImageToDocxConverter(output_path, filename, image_size, margin_mm)
    converter.run(image_paths)


if __name__ == "__main__":
    main(sys.argv[1:])  # Exclude the script name from arguments
