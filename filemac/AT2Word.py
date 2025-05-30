"""Create a word document directly from a text file."""

from docx import Document
from docx.shared import Pt, RGBColor

from utils.colors import foreground

fcl = foreground()
RESET = fcl.RESET


class AdvancedT2word:
    """
    Args:
        obj-> input object (normally a formated text file)
        fsize ->font-size default = 12: int
        fstyle -> font-name default = Times New Roman: str
        out_obj -> output object(file) name: str
    Returns:
        None

        Given obj -> Text file where:
        '#' is used to specify formarting
        Only three heading leavels are supported.
        '#' Heading1,
        '##' -> Heading2,
        '###' -> Heading3
    """

    def __init__(
        self, obj, out_obj=None, fsize: int = 12, fstyle: str = "Times New Roman"
    ):
        self.obj = obj
        self.out_obj = out_obj
        self.fsize = fsize
        self.fstyle = fstyle
        if self.out_obj is None:
            self.out_obj = f"{self.obj.split('.')[0]}_filemac.docx"

    def text_to_word(self):
        """
        Create new document,
        heading_styles -> define formating
        Open the text file and read it line by line.
        For every line check whether it starts with '#' format specify , ommit the specifier and formart the line.
        Strip empty spaces from every line.
        Set body font to fstyle and font size to fsize.
        """

        print(f"{fcl.BWHITE_FG}Set Font: {fcl.CYAN_FG}{self.fsize}{RESET}")
        print(f"{fcl.BWHITE_FG}Set Style: {fcl.CYAN_FG}{self.fstyle}{RESET}")
        # Create a new Document
        doc = Document()

        # Define formatting for headings and body text
        head_font_name = self.fstyle
        heading_styles = {
            # Heading 1
            1: {"font_size": Pt(18), "font_color": RGBColor(126, 153, 184)},
            # Heading 2
            2: {"font_size": Pt(16), "font_color": RGBColor(0, 120, 212)},
            # Heading 3
            3: {"font_size": Pt(14), "font_color": RGBColor(0, 120, 212)},
            # Heading 4
            4: {"font_size": Pt(13), "font_color": RGBColor(0, 120, 212)},
        }

        body_font_name = "Times New Roman"
        body_font_size = Pt(self.fsize)
        body_font_color = RGBColor(0, 0, 0)  # Black color

        # Open the text file and read content
        with open(self.obj, "r") as file:
            lines = file.readlines()

        for i, line in enumerate(lines):
            print(
                f"{fcl.BWHITE_FG}Line: {fcl.DCYAN_FG}{i}{fcl.YELLOW_FG} of {fcl.BLUE_FG}{len(lines)}{RESET}",
                end="\r",
            )
            # Determine heading level or body text
            if line.startswith("#"):
                level = line.count("#")
                level = min(level, 3)  # Support up to 3 levels of headings
                style = heading_styles.get(level, heading_styles[1])
                p = doc.add_paragraph()
                # Remove '#' and extra space
                run = p.add_run(line[level + 1 :].strip())
                run.font.size = style["font_size"]
                run.font.name = head_font_name
                run.font.color.rgb = style["font_color"]
                p.style = f"Heading{level}"
            else:
                p = doc.add_paragraph()
                run = p.add_run(line.strip())
                run.font.name = body_font_name
                run.font.size = body_font_size
                run.font.color.rgb = body_font_color

        # Save the document
        print("\n")
        doc.save(self.out_obj)
        print(
            f"{fcl.BWHITE_FG}Text file converted to Word document: {fcl.MAGENTA_FG}{self.out_obj}{RESET}"
        )


if __name__ == "__main__":
    init = AdvancedT2word("/home/skye/Documents/FMAC/file2.txt")

    # Call the function
    init.text_to_word()
