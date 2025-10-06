"""
Convert Markdown to DOCX with Mermaid fenced blocks rendered into diagrams.

Usage:
    converter = MarkdownToDocxConverter()
    converter.convert_file("example.md", "output.docx")
"""

import os
import io
import subprocess
import tempfile
import logging
from typing import Optional, List, Tuple
from pathlib import Path
from urllib.parse import quote_plus

import requests
from markdown_it import MarkdownIt
from markdown_it.token import Token
from docx import Document
from docx.shared import Inches, Pt
from ..exceptions.handler import MermaidRenderError

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class MermaidRenderer:
    """
    Renders mermaid diagram source to an image (PNG) either via the local `mmdc` CLI
    or via the mermaid.ink web service as a fallback.
    """

    def __init__(self, mmdc_path: str = "mmdc", timeout: int = 15):
        self.mmdc_path = mmdc_path
        self.timeout = timeout

    def render_to_png(self, mermaid_source: str) -> bytes:
        """
        Returns PNG bytes.
        Try local mmdc first; fallback to mermaid.ink.
        """
        try:
            logger.debug("Trying local mmdc renderer")
            return self._render_with_mmdc(mermaid_source)
        except Exception as e:
            logger.info("Local mmdc failed: %s. Falling back to mermaid.ink", e)
            try:
                return self._render_with_mermaid_ink(mermaid_source)
            except Exception as e2:
                logger.error("mermaid.ink fallback failed: %s", e2)
                raise MermaidRenderError(
                    f"Failed to render mermaid diagram: {e2}"
                ) from e2

    def _render_with_mmdc(self, source: str) -> bytes:
        """
        Use the mermaid-cli (mmdc) to render PNG. Requires Node.js + @mermaid-js/mermaid-cli installed.
        """
        with tempfile.NamedTemporaryFile(
            suffix=".mmd", mode="w+", delete=False
        ) as srcf:
            srcf.write(source)
            srcf.flush()
            src_name = srcf.name

        out_fd, out_name = tempfile.mkstemp(suffix=".png")
        os.close(out_fd)

        cmd = [
            self.mmdc_path,
            "-i",
            src_name,
            "-o",
            out_name,
            "-b",
            "transparent",  # background transparent if supported
            # optionally: "-w", "800", "-H", "600"
        ]
        logger.debug("Running command: %s", " ".join(cmd))
        try:
            subprocess.run(
                cmd,
                check=True,
                timeout=self.timeout,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            with open(out_name, "rb") as f:
                png_bytes = f.read()
            return png_bytes
        finally:
            try:
                os.remove(src_name)
            except OSError:
                pass
            try:
                os.remove(out_name)
            except OSError:
                pass

    def _render_with_mermaid_ink(self, source: str, image_format: str = "png") -> bytes:
        """
        Use mermaid.ink service to get the diagram image.
        It expects the diagram text URL-encoded as part of the path.
        Example: https://mermaid.ink/img/png/<urlencoded source>
        """
        base = "https://mermaid.ink/img"
        # URL-encode the mermaid source. mermaid.ink accepts encodeURIComponent style.
        encoded = quote_plus(source)  # safe fallback
        url = f"{base}/{image_format}/{encoded}"
        logger.debug("Requesting mermaid.ink URL: %s", url)
        resp = requests.get(url, timeout=self.timeout)
        if resp.status_code != 200:
            raise MermaidRenderError(
                f"mermaid.ink returned status {resp.status_code}: {resp.text[:200]}"
            )
        return resp.content


class MarkdownToDocxConverter:
    """
    Convert Markdown text to a DOCX Document, rendering Mermaid blocks as embedded images.
    """

    def __init__(
        self,
        mermaid_renderer: Optional[MermaidRenderer] = None,
        default_font: str = "Calibri",
        default_font_size: int = 11,
    ):
        self.md = MarkdownIt("commonmark")
        self.mermaid_renderer = mermaid_renderer or MermaidRenderer()
        self.document = Document()
        # Set normal style font
        style = self.document.styles["Normal"]
        style.font.name = default_font
        style.font.size = Pt(default_font_size)

    def convert_text(self, markdown_text: str) -> Document:
        """
        Parse markdown_text and populate self.document.
        Returns the python-docx Document object.
        """
        tokens = self.md.parse(markdown_text)
        # We'll walk tokens and dispatch handling
        i = 0
        while i < len(tokens):
            tok = tokens[i]
            logger.debug("Token: %s %s", tok.type, tok.tag)
            if tok.type == "heading_open":
                level = int(tok.tag[1])
                # next token should be inline with children
                inline = tokens[i + 1]
                text = self._inline_to_text(inline)
                self._add_heading(text, level)
                i += 3  # skip heading_open, inline, heading_close
                continue
            if tok.type == "fence":
                lang = tok.info.strip().split()[0] if tok.info else ""
                content = tok.content
                if lang.lower() in ("mermaid", "mmd"):
                    self._handle_mermaid_block(content)
                else:
                    self._add_code_block(content, lang)
                i += 1
                continue
            if tok.type == "paragraph_open":
                # next = inline, following = paragraph_close
                inline = tokens[i + 1]
                text_runs = self._inline_to_runs(inline)
                self._add_paragraph_with_runs(text_runs)
                i += 3
                continue
            if tok.type in ("bullet_list_open", "ordered_list_open"):
                i = self._handle_list(tokens, i)
                continue
            # ignore other tokens for now (blockquote, html_block, etc.)
            i += 1

        return self.document

    def convert_file(self, md_path: str, docx_path: str) -> Path:
        md_text = Path(md_path).read_text(encoding="utf-8")
        doc = self.convert_text(md_text)
        doc.save(docx_path)
        logger.info("Saved DOCX: %s", docx_path)

        return docx_path

    # ---------- Handlers for specific blocks ----------
    def _add_heading(self, text: str, level: int):
        # docx levels: 0=Title, 1=Heading1,...
        level = max(
            0, min(4, level - 1)
        )  # map 1->0 or use docx add_heading level directly
        # python-docx add_heading uses levels 0..4 -> Heading 1..5?
        self.document.add_heading(text, level=level)
        logger.debug("Added heading level %s: %s", level, text)

    def _add_code_block(self, code: str, language: str = ""):
        p = self.document.add_paragraph()
        run = p.add_run(code.rstrip("\n"))
        # monospace style
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        # Optionally add shading/border - python-docx lacks easy paragraph background; skip for brevity.
        logger.debug("Added code block (lang=%s)", language)

    def _handle_mermaid_block(self, source: str):
        try:
            png_bytes = self.mermaid_renderer.render_to_png(source)
        except MermaidRenderError as e:
            logger.error("Failed to render Mermaid diagram: %s", e)
            # Insert source as preformatted text as a fallback
            self._add_code_block(source, language="mermaid (failed)")
            return

        # Insert image into docx
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as imgf:
            imgf.write(png_bytes)
            imgf.flush()
            img_path = imgf.name

        try:
            # Choose a reasonable width (max 6 inches). python-docx scales automatically.
            self.document.add_picture(img_path, width=Inches(6))
            logger.debug("Embedded mermaid image into docx")
        finally:
            try:
                os.remove(img_path)
            except OSError:
                pass

    # ---------- Inline parsing ----------
    def _inline_to_text(self, inline_token: Token) -> str:
        """
        Simple flattening of inline token to plain text.
        """
        if inline_token.type != "inline":
            return ""
        parts = []
        for child in inline_token.children or []:
            parts.append(child.content or "")
        return "".join(parts)

    def _inline_to_runs(
        self, inline_token: Token, active_styles=None
    ) -> List[Tuple[str, dict]]:
        """
        Recursively converts inline tokens into a list of (text, style_dict) tuples.
        Supports bold, italic, code, and hyperlinks.
        """
        runs = []
        if active_styles is None:
            active_styles = {}

        if inline_token.type != "inline":
            return runs

        children = inline_token.children or []
        i = 0
        while i < len(children):
            child = children[i]
            ttype = child.type

            # TEXT
            if ttype == "text":
                runs.append((child.content, active_styles.copy()))

            # SOFT BREAK / LINE BREAK
            elif ttype in ("softbreak", "hardbreak"):
                runs.append(("\n", active_styles.copy()))

            # INLINE CODE
            elif ttype == "code_inline":
                style = active_styles.copy()
                style.update({"code": True})
                runs.append((child.content, style))

            # BOLD
            elif ttype == "strong_open":
                # collect until strong_close
                sub_tokens = self._collect_until(children, i + 1, "strong_close")
                sub_inline = Token("inline", "", 0)
                sub_inline.children = sub_tokens
                style = active_styles.copy()
                style.update({"bold": True})
                runs.extend(self._inline_to_runs(sub_inline, style))
                # move index to after strong_close
                i = self._find_close(children, i + 1, "strong_close")

            # ITALIC
            elif ttype == "em_open":
                sub_tokens = self._collect_until(children, i + 1, "em_close")
                sub_inline = Token("inline", "", 0)
                sub_inline.children = sub_tokens
                style = active_styles.copy()
                style.update({"italic": True})
                runs.extend(self._inline_to_runs(sub_inline, style))
                i = self._find_close(children, i + 1, "em_close")

            # LINK
            elif ttype == "link_open":
                href = dict(child.attrs or {}).get("href", "")
                sub_tokens = self._collect_until(children, i + 1, "link_close")
                sub_inline = Token("inline", "", 0)
                sub_inline.children = sub_tokens
                style = active_styles.copy()
                style.update({"link": href})
                runs.extend(self._inline_to_runs(sub_inline, style))
                i = self._find_close(children, i + 1, "link_close")

            i += 1

        return runs

    def _collect_until(self, tokens, start_index, close_type):
        """Collect tokens until close_type (exclusive)."""
        collected = []
        i = start_index
        while i < len(tokens) and tokens[i].type != close_type:
            collected.append(tokens[i])
            i += 1
        return collected

    def _find_close(self, tokens, start_index, close_type):
        """Find index of closing token; return its index (will be incremented in main loop)."""
        i = start_index
        while i < len(tokens):
            if tokens[i].type == close_type:
                return i
            i += 1
        return len(tokens) - 1

    def _add_paragraph_with_runs(self, text_runs: List[Tuple[str, dict]]):
        p = self.document.add_paragraph()
        for txt, style in text_runs:
            r = p.add_run(txt)
            if style.get("bold", None):
                r.bold = True
            if style.get("italic", None):
                r.italic = True
            if style.get("code", None):
                r.font.name = "Consolas"
                r.font.size = Pt(9)
            if style.get("link", None):
                # Python-docx doesn’t natively support clickable hyperlinks, so we underline + color blue
                r.font.color.rgb = (0x00, 0x00, 0xFF)
                r.underline = True
        logger.debug("Added paragraph with styled runs")

    def _add_paragraph_with_runs_dummy(self, text_runs: List[Tuple[str, dict]]):
        p = self.document.add_paragraph()
        for txt, style in text_runs:
            r = p.add_run(txt)
            if style.get("code"):
                r.font.name = "Consolas"
                r.font.size = Pt(9)
        logger.debug("Added paragraph with runs")

    # ---------- List handling ----------
    def _handle_list(self, tokens: List[Token], start_index: int) -> int:
        """
        Parse a list starting at start_index. Return new index after the list.
        Basic implementation: iterate until list_close and add paragraphs with bullet or numbering.
        """
        list_tok = tokens[start_index]
        is_ordered = list_tok.type == "ordered_list_open"
        i = start_index + 1
        while i < len(tokens) and tokens[i].type != (
            "ordered_list_close" if is_ordered else "bullet_list_close"
        ):
            tok = tokens[i]
            if tok.type == "list_item_open":
                # Expect paragraph_open -> inline -> paragraph_close -> list_item_close
                if (
                    tokens[i + 1].type == "paragraph_open"
                    and tokens[i + 2].type == "inline"
                ):
                    inline = tokens[i + 2]
                    text = self._inline_to_text(inline)
                    p = self.document.add_paragraph(text)
                    if is_ordered:
                        p.style = "List Number"
                    else:
                        p.style = "List Bullet"
                    # list_item_open, paragraph_open, inline, paragraph_close (then list_item_close)
                    i += 4
                    # there might be nested items; for brevity skip nested lists support here
                    continue
                else:
                    i += 1
                    continue
            else:
                i += 1
        # now tokens[i] is list_close
        return i + 1


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="Convert Markdown to DOCX with Mermaid rendering."
    )
    parser.add_argument("input", help="Input markdown file")
    parser.add_argument("output", help="Output docx file")
    args = parser.parse_args()

    renderer = MermaidRenderer()
    converter = MarkdownToDocxConverter(mermaid_renderer=renderer)
    converter.convert_file(args.input, args.output)
    print("Done.")
"""
# Python packages
pip install python-docx markdown-it-py requests

# Optional (recommended) for best local mermaid rendering:
# Node.js + mermaid-cli
npm install -g @mermaid-js/mermaid-cli

# (If you want nicer code highlighting later, you can add Pygments and render code snippets to images.)
"""
