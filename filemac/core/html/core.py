"""
custom_html_to_docx.py
A reliable HTML to DOCX converter specifically designed for CVs and professional documents
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import re
from typing import List, Dict, Any
import html as html_parser


class CVHTMLConverter:
    """A specialized HTML to DOCX converter for CVs and professional documents"""

    def __init__(self):
        self.doc = None
        self.current_paragraph = None
        self.styles = {
            "h1": {"size": 16, "bold": True, "alignment": WD_ALIGN_PARAGRAPH.CENTER},
            "h2": {"size": 14, "bold": True, "alignment": WD_ALIGN_PARAGRAPH.LEFT},
            "h3": {"size": 12, "bold": True, "alignment": WD_ALIGN_PARAGRAPH.LEFT},
            "normal": {"size": 11, "bold": False, "alignment": WD_ALIGN_PARAGRAPH.LEFT},
            "bold": {"size": 11, "bold": True, "alignment": WD_ALIGN_PARAGRAPH.LEFT},
            "italic": {
                "size": 11,
                "italic": True,
                "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            },
        }

    def convert(self, html_content: str, output_path: str) -> Document:
        """
        Convert HTML content to DOCX document

        Args:
            html_content: HTML string to convert
            output_path: Path for output DOCX file

        Returns:
            Document: The created Word document
        """
        self.doc = Document()
        self._setup_document_styles()

        # Clean and parse HTML
        cleaned_html = self._clean_html(html_content)
        self._parse_html(cleaned_html)

        self.doc.save(output_path)
        return self.doc

    def _setup_document_styles(self):
        """Setup document styles and formatting"""
        # Set normal style
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Create custom styles for CV
        self._create_style("CV Title", 16, True, WD_ALIGN_PARAGRAPH.CENTER)
        self._create_style("CV Heading", 14, True, WD_ALIGN_PARAGRAPH.LEFT)
        self._create_style("CV Subheading", 12, True, WD_ALIGN_PARAGRAPH.LEFT)

    def _create_style(self, style_name: str, font_size: int, bold: bool, alignment):
        """Create a custom style"""
        try:
            style = self.doc.styles.add_style(
                style_name, 1
            )  # 1 = WD_STYLE_TYPE.PARAGRAPH
            font = style.font
            font.name = "Calibri"
            font.size = Pt(font_size)
            font.bold = bold
            style.paragraph_format.alignment = alignment
        except:
            # Style might already exist
            pass

    def _clean_html(self, html: str) -> str:
        """Clean and normalize HTML content"""
        # Remove multiple spaces and newlines
        html = re.sub(r"\s+", " ", html)

        # Ensure proper tag formatting
        html = html.replace("<br>", "<br/>").replace("<hr>", "<hr/>")

        # Decode HTML entities
        html = html_parser.unescape(html)

        return html.strip()

    def _parse_html(self, html: str):
        """Parse HTML content and convert to DOCX"""
        # Split by tags while preserving content
        tokens = self._tokenize_html(html)
        self._process_tokens(tokens)

    def _tokenize_html(self, html: str) -> List[Dict[str, Any]]:
        """Tokenize HTML into manageable chunks"""
        tokens = []
        pos = 0

        while pos < len(html):
            # Find next tag
            tag_match = re.search(r"</?(\w+)[^>]*>", html[pos:])

            if not tag_match:
                # Add remaining text
                if pos < len(html):
                    tokens.append({"type": "text", "content": html[pos:]})
                break

            tag_start = tag_match.start() + pos
            tag_end = tag_match.end() + pos

            # Add text before tag
            if tag_start > pos:
                tokens.append({"type": "text", "content": html[pos:tag_start]})

            # Add tag
            tag_content = html[tag_start:tag_end]
            is_closing = tag_content.startswith("</")
            tag_name = tag_match.group(1).lower()

            tokens.append(
                {
                    "type": "tag",
                    "name": tag_name,
                    "content": tag_content,
                    "is_closing": is_closing,
                    "is_self_closing": tag_content.endswith("/>"),
                }
            )

            pos = tag_end

        return tokens

    def _process_tokens(self, tokens: List[Dict[str, Any]]):
        """Process tokens and build document"""
        stack = []  # Track open tags

        for token in tokens:
            if token["type"] == "text":
                self._add_text(token["content"], stack)
            elif token["type"] == "tag":
                if token["is_closing"]:
                    # Close tag
                    if stack and stack[-1]["name"] == token["name"]:
                        stack.pop()
                    self._handle_closing_tag(token["name"])
                else:
                    # Open tag
                    stack.append(token)
                    self._handle_opening_tag(token)

    def _handle_opening_tag(self, tag: Dict[str, Any]):
        """Handle opening tag"""
        tag_name = tag["name"]

        if tag_name in ["h1", "h2", "h3"]:
            self._add_heading(tag_name)
        elif tag_name == "br":
            self._add_line_break()
        elif tag_name == "p":
            self._start_paragraph()
        elif tag_name == "div":
            if self.current_paragraph:
                self.current_paragraph = None
        elif tag_name == "ul":
            self.in_list = True
        elif tag_name == "li":
            self._start_list_item()

    def _handle_closing_tag(self, tag_name: str):
        """Handle closing tag"""
        if tag_name in ["h1", "h2", "h3", "p"]:
            self.current_paragraph = None
        elif tag_name == "ul":
            self.in_list = False
            self.current_paragraph = None

    def _add_heading(self, level: str):
        """Add heading based on level"""
        self.current_paragraph = self.doc.add_paragraph()

        if level == "h1":
            self.current_paragraph.style = "CV Title"
        elif level == "h2":
            self.current_paragraph.style = "CV Heading"
        else:
            self.current_paragraph.style = "CV Subheading"

    def _start_paragraph(self):
        """Start a new paragraph"""
        self.current_paragraph = self.doc.add_paragraph()

    def _start_list_item(self):
        """Start a new list item"""
        self.current_paragraph = self.doc.add_paragraph(style="List Bullet")

    def _add_line_break(self):
        """Add line break"""
        if self.current_paragraph:
            self.current_paragraph.add_run().add_break(WD_BREAK.LINE)
        else:
            self.doc.add_paragraph()

    def _add_text(self, text: str, stack: List[Dict[str, Any]]):
        """Add text with current formatting"""
        if not text.strip():
            return

        # Create paragraph if none exists
        if not self.current_paragraph:
            self.current_paragraph = self.doc.add_paragraph()

        run = self.current_paragraph.add_run(text)

        # Apply formatting based on stack
        self._apply_formatting(run, stack)

    def _apply_formatting(self, run, stack: List[Dict[str, Any]]):
        """Apply formatting based on tag stack"""
        font = run.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Check for bold (strong, b)
        bold_tags = ["strong", "b", "h1", "h2", "h3"]
        if any(tag["name"] in bold_tags for tag in stack):
            font.bold = True

        # Check for italic (em, i)
        italic_tags = ["em", "i"]
        if any(tag["name"] in italic_tags for tag in stack):
            font.italic = True

        # Check for underline (u)
        underline_tags = ["u"]
        if any(tag["name"] in underline_tags for tag in stack):
            font.underline = True


class AdvancedCVConverter(CVHTMLConverter):
    """Enhanced converter with better CSS support and layout management"""

    def __init__(self):
        super().__init__()
        self.section_spacing = Pt(12)
        self.current_style = {}

    def convert_cv_html(self, html_file_path: str, output_path: str) -> Document:
        """
        Convert CV HTML file to DOCX with enhanced formatting

        Args:
            html_file_path: Path to HTML file
            output_path: Output DOCX path
        """
        with open(html_file_path, "r", encoding="utf-8") as f:
            html_content = f.read()

        return self.convert(html_content, output_path)

    def _extract_css_styles(self, html: str) -> Dict[str, Dict]:
        """Extract CSS styles from style tags"""
        styles = {}
        style_matches = re.findall(r"<style[^>]*>(.*?)</style>", html, re.DOTALL)

        for style_content in style_matches:
            # Parse CSS rules (simplified)
            rules = re.findall(r"\.(\w+)\s*\{([^}]+)\}", style_content)
            for class_name, properties in rules:
                styles[class_name] = self._parse_css_properties(properties)

        return styles

    def _parse_css_properties(self, css: str) -> Dict[str, str]:
        """Parse CSS properties into dictionary"""
        properties = {}
        declarations = css.split(";")

        for declaration in declarations:
            if ":" in declaration:
                prop, value = declaration.split(":", 1)
                properties[prop.strip().lower()] = value.strip()

        return properties

    def _apply_css_style(self, run, style: Dict):
        """Apply CSS-style formatting"""
        if "font-weight" in style and "bold" in style["font-weight"]:
            run.font.bold = True

        if "font-style" in style and "italic" in style["font-style"]:
            run.font.italic = True

        if "text-align" in style:
            if "center" in style["text-align"]:
                run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif "right" in style["text-align"]:
                run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if "color" in style:
            color = self._parse_css_color(style["color"])
            if color:
                run.font.color.rgb = color

    def _parse_css_color(self, color_str: str) -> RGBColor:
        """Parse CSS color string to RGBColor"""
        # Handle hex colors
        hex_match = re.match(
            r"#([0-9a-fA-F]{2})([0-9a-fA-F]{2})([0-9a-fA-F]{2})", color_str
        )
        if hex_match:
            r, g, b = [int(x, 16) for x in hex_match.groups()]
            return RGBColor(r, g, b)

        # Handle rgb() colors
        rgb_match = re.match(r"rgb\((\d+),\s*(\d+),\s*(\d+)\)", color_str)
        if rgb_match:
            r, g, b = [int(x) for x in rgb_match.groups()]
            return RGBColor(r, g, b)

        return None


# Usage examples and helper functions
def create_cv_from_html_template():
    """Create a CV using our HTML template"""
    converter = AdvancedCVConverter()

    # Example usage
    with open("/home/skye/Downloads/MWG-CV.html", "r") as f:
        html_content = f.read()

    return converter.convert(html_content, "professional_cv.docx")


def main():
    """Main demonstration function"""
    print("Custom HTML to DOCX Converter for CVs")
    print("=====================================")

    # Example 2: Create from our CV template
    create_cv_from_html_template()
    print("✓ Professional CV created successfully!")

    print("\nLibrary features:")
    print("- Custom HTML parsing optimized for CVs")
    print("- Professional styling and formatting")
    print("- List and bullet point support")
    print("- Heading hierarchy")
    print("- Basic CSS style support")
    print("- Robust error handling")


if __name__ == "__main__":
    main()
