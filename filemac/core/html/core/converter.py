"""
Main converter class that orchestrates the HTML to DOCX conversion
"""

import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches
from typing import Dict, List
from docx.oxml import OxmlElement
from .html_parser import HTMLParser
from .style_manager import StyleManager
from ..utils.validation import validate_html, validate_file_path


class HTML2Word:
    """Main converter class that coordinates HTML parsing and DOCX generation"""

    def __init__(self, default_font: str = "Calibri", default_size: int = 11):
        self.doc = None
        self.default_font = default_font
        self.default_size = default_size
        self.html_parser = HTMLParser()
        self.style_manager = StyleManager(default_font, default_size)

        # Conversion state
        self.current_paragraph = None
        self.current_style = {}
        self.tag_stack = []
        self.block_element = {}

    def __enter__(self):
        self.block_element = self.html_parser.block_elements.copy().add("ol").add("ul")

    def convert(self, html_content: str, output_path: str) -> Document:
        """
        Convert HTML content to DOCX document

        Args:
            html_content: HTML string to convert
            output_path: Path for output DOCX file

        Returns:
            Document: The created Word document
        """
        # Validate inputs
        validate_html(html_content)
        validate_file_path(output_path, "output")

        # Initialize document
        self.doc = Document()
        self.style_manager.setup_document_styles(self.doc)

        # Parse HTML and extract styles
        parsed_data = self.html_parser.parse(html_content)

        # Convert to DOCX
        self._convert_elements(parsed_data["elements"], parsed_data["styles"])

        # Save document
        self.doc.save(output_path)
        return self.doc

    def convert_file(self, html_file_path: str, output_path: str) -> Document:
        """
        Convert HTML file to DOCX document

        Args:
            html_file_path: Path to HTML file
            output_path: Path for output DOCX file

        Returns:
            Document: The created Word document
        """
        validate_file_path(html_file_path, "input")

        with open(html_file_path, "r", encoding="utf-8") as f:
            html_content = f.read()

        return self.convert(html_content, output_path)

    def _convert_elements(self, elements: List[Dict], styles: Dict):
        """Convert parsed HTML elements to DOCX format"""
        for element in elements:
            self._convert_element(element, styles)

    def _convert_element(self, element: Dict, styles: Dict):
        """Convert a single HTML element to DOCX"""
        element_type = element["type"]

        if element_type == "text":
            self._add_text_element(element, styles)
        elif element_type == "element":
            self._handle_html_element(element, styles)

    def _handle_html_element(self, element: Dict, styles: Dict):
        """Handle HTML element based on tag type"""
        tag_name = element["tag"].lower()

        # Push to stack for styling
        self.tag_stack.append(element)

        is_block_element = tag_name in self.block_element

        try:
            # For block elements, ensure we start a new paragraph context
            if is_block_element and self.current_paragraph:
                # Only start new paragraph if the current one has content
                if self.current_paragraph.text.strip():
                    self.current_paragraph = None

            # Check for grid/flex containers first
            if tag_name in ("div", "section", "container"):
                attributes = element.get("attributes", {})
                style_attr = attributes.get("style", "")
                is_grid = "display:grid" in style_attr.replace(" ", "")
                is_flex = "display:flex" in style_attr.replace(" ", "")

                if is_grid or is_flex:
                    self._handle_grid_container(element, styles)
                else:
                    self._add_div(element, styles)

            elif tag_name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                self._add_heading(element, styles)
            elif tag_name == "p":
                self._add_paragraph(element, styles)
            elif tag_name == "table":
                self._add_table(element, styles)
            elif tag_name == "span":
                self._add_span(element, styles)
            elif tag_name == "br":
                self._add_line_break()
            elif tag_name == "hr":
                self._add_horizontal_rule()
            elif tag_name == "ul":
                self._start_list(element, styles)
            elif tag_name == "ol":
                self._start_numbered_list(element, styles)
            elif tag_name == "li":
                self._add_list_item(element, styles)
            elif tag_name == "strong" or tag_name == "b":
                self._add_bold_text(element, styles)
            elif tag_name == "em" or tag_name == "i":
                self._add_italic_text(element, styles)
            elif tag_name == "u":
                self._add_underline_text(element, styles)
            elif tag_name == "pre":
                self._add_preformatted_text(element, styles)
            elif tag_name == "tr":
                self._add_table_row(element, styles)
            elif tag_name == "td" or tag_name == "th":
                self._add_table_cell(element, styles)
            else:
                # Process children for unknown tags
                self._convert_elements(element.get("children", []), styles)

        finally:
            # Always pop from stack
            if self.tag_stack and self.tag_stack[-1] == element:
                self.tag_stack.pop()

        # For block elements, ensure we clear the paragraph context after processing
        if is_block_element:
            self.current_paragraph = None

    def _add_preformatted_text(self, element: Dict, styles: Dict):
        """Add preformatted text with preserved whitespace"""
        if not self.current_paragraph:
            self.current_paragraph = self.doc.add_paragraph()

        # Process all text content in pre tag
        self._process_preformatted_content(element.get("children", []), styles)
        self.current_paragraph = None

    def _process_preformatted_content(self, elements: List[Dict], styles: Dict):
        """Process content for pre tags with preserved formatting"""
        for element in elements:
            if element["type"] == "text":
                text = element.get("content", "")
                if text:
                    # Preserve all whitespace in pre tags
                    run = self.current_paragraph.add_run(text)
                    self.style_manager.apply_styles_to_run(run, self.tag_stack, styles)
            elif element["type"] == "element":
                self._handle_html_element(element, styles)

    def _add_text_element(self, element: Dict, styles: Dict):
        """Add text element with proper styling and line breaks"""
        if not self.current_paragraph:
            self.current_paragraph = self.doc.add_paragraph()

        text = element.get("content", "")

        # Handle text with line breaks
        if "\n" in text:
            lines = text.split("\n")
            for i, line in enumerate(lines):
                if line.strip():  # Only add non-empty lines
                    run = self.current_paragraph.add_run(line.strip())
                    self.style_manager.apply_styles_to_run(run, self.tag_stack, styles)

                # Add line break except for the last line
                if i < len(lines) - 1 and line.strip():
                    self._add_line_break()
        else:
            # Single line of text
            if text.strip():
                run = self.current_paragraph.add_run(text.strip())
                self.style_manager.apply_styles_to_run(run, self.tag_stack, styles)

    def _add_heading(self, element: Dict, styles: Dict):
        """Add heading with appropriate level"""
        level = element["tag"][1]  # Extract number from h1, h2, etc.
        self.current_paragraph = self.doc.add_paragraph()

        # Apply heading style
        self.style_manager.apply_heading_style(
            self.current_paragraph, level, element, styles
        )

        # Process children
        self._convert_elements(element.get("children", []), styles)

        self.current_paragraph = None

    def _add_paragraph(self, element: Dict, styles: Dict):
        """Add paragraph"""
        self.current_paragraph = self.doc.add_paragraph()
        self.style_manager.apply_paragraph_style(
            self.current_paragraph, element, styles
        )

        # Process children
        self._convert_elements(element.get("children", []), styles)

        self.current_paragraph = None

    def _add_div(self, element: Dict, styles: Dict):
        """Add div element - ensure new paragraph for block-level elements"""
        # For divs that contain block-level content, start a new paragraph
        attributes = element.get("attributes", {})
        display_style = attributes.get("style", "")
        try:
            # Check for display: inline in styles
            is_inline = "display:inline" in display_style
        except AttributeError:
            pass

        has_block_content = (
            any(
                child.get("tag") in self.block_element
                for child in element.get("children", [])
                if child["type"] == "element"
            )
            or is_inline
        )

        if has_block_content and self.current_paragraph:
            self.current_paragraph = None

        self._convert_elements(element.get("children", []), styles)

    def _add_span(self, element: Dict, styles: Dict):
        """Add span with inline styling"""
        self._convert_elements(element.get("children", []), styles)

    def _add_line_break(self):
        """Add a proper line break in Word document"""
        if self.current_paragraph:
            # Only add break if the paragraph has content
            if self.current_paragraph.text.strip():
                self.current_paragraph.add_run().add_break()
            else:
                # If empty, add a space to maintain the break
                self.current_paragraph.add_run(" ")
        else:
            # Create a new paragraph for the line break
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.add_run(" ")

        # def _add_horizontal_rule(self):
        """Add horizontal rule"""
        # self.doc.add_paragraph().add_run("_" * 50)

    def _add_horizontal_rule(self):
        """Add a proper horizontal rule/line"""
        try:
            # Create a new paragraph for the horizontal rule
            hr_paragraph = self.doc.add_paragraph()

            # Add border to the paragraph to create the horizontal line
            p_pr = hr_paragraph._p.get_or_add_pPr()

            # Create paragraph borders
            p_borders = OxmlElement("w:pBdr")

            # Create bottom border for the horizontal line
            bottom_border = OxmlElement("w:bottom")
            bottom_border.set(qn("w:val"), "single")
            bottom_border.set(qn("w:sz"), "6")  # Line thickness (6 = 0.75 pt)
            bottom_border.set(qn("w:space"), "1")  # Spacing above the line
            bottom_border.set(qn("w:color"), "auto")  # Automatic color

            # Add the bottom border to the borders element
            p_borders.append(bottom_border)

            # Add borders to paragraph properties
            p_pr.append(p_borders)

            # Add some spacing after the horizontal rule
            p_spacing = OxmlElement("w:spacing")
            p_spacing.set(qn("w:after"), "120")  # 120 twips = 6 points spacing after
            p_pr.append(p_spacing)

            # Clear current paragraph context
            self.current_paragraph = None

        except Exception as e:
            # Fallback: create a simple horizontal line with underscores
            fallback_paragraph = self.doc.add_paragraph()
            fallback_paragraph.add_run("_" * 50)  # Simple underscore line
            self.current_paragraph = None
            print(f"Horizontal rule fallback used: {e}")

    def _start_list(self, element: Dict, styles: Dict):
        """Start unordered list"""
        self._convert_elements(element.get("children", []), styles)

    def _start_numbered_list(self, element: Dict, styles: Dict):
        """Start ordered list"""
        self._convert_elements(element.get("children", []), styles)

    def _add_list_item(self, element: Dict, styles: Dict):
        """Add list item"""
        self.current_paragraph = self.doc.add_paragraph(style="List Bullet")
        self._convert_elements(element.get("children", []), styles)
        self.current_paragraph = None

    def _add_bold_text(self, element: Dict, styles: Dict):
        """Add bold text"""
        self._convert_elements(element.get("children", []), styles)

    def _add_italic_text(self, element: Dict, styles: Dict):
        """Add italic text"""
        self._convert_elements(element.get("children", []), styles)

    def _add_underline_text(self, element: Dict, styles: Dict):
        """Add underline text"""
        self._convert_elements(element.get("children", []), styles)

    # Table handling methods
    def _add_table(self, element: Dict, styles: Dict):
        """Create a new table"""
        # Save current state
        saved_paragraph = self.current_paragraph

        attributes = element.get("attributes", {})

        # Calculate table dimensions
        rows = self._count_table_rows(element)
        cols = self._count_table_columns(element)

        try:
            # Create table with calculated dimensions
            if rows > 0 and cols > 0:
                # Create table at current position
                self.current_table = self.doc.add_table(rows=rows, cols=cols)
                self.current_table.autofit = True
                self.current_table.allow_autofit = True

                # Apply table styles
                self._apply_table_styles(self.current_table, attributes, styles)

                # Reset row and cell counters
                self.current_row_index = 0
                self.current_cell_index = 0

                # Process table content
                self._convert_elements(element.get("children", []), styles)

                # Clean up empty rows/cells if needed
                self._cleanup_table()

            # Add a paragraph after the table for proper flow
            self.current_paragraph = self.doc.add_paragraph()

        except Exception as e:
            print(f"Table creation error: {e}")
            # Fallback: add a simple paragraph
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.add_run("[Table content]")

        finally:
            # Restore or clear table context
            self.current_table = None
            # Restore paragraph context if it was saved
            if saved_paragraph:
                self.current_paragraph = saved_paragraph

    def _count_table_rows(self, table_element: Dict) -> int:
        """Count the number of rows in the table"""
        rows = 0
        for child in table_element.get("children", []):
            if child.get("type") == "element" and child.get("tag", "").lower() == "tr":
                rows += 1
        return max(rows, 1)  # At least 1 row

    def _count_table_columns(self, table_element: Dict) -> int:
        """Count the maximum number of columns in the table"""
        max_cols = 0
        for child in table_element.get("children", []):
            if child.get("type") == "element" and child.get("tag", "").lower() == "tr":
                col_count = 0
                for cell in child.get("children", []):
                    if cell.get("type") == "element" and cell.get(
                        "tag", ""
                    ).lower() in [
                        "td",
                        "th",
                    ]:
                        # Handle colspan
                        colspan = int(cell.get("attributes", {}).get("colspan", 1))
                        col_count += colspan
                max_cols = max(max_cols, col_count)
        return max(max_cols, 1)  # At least 1 column

    def _add_table_row(self, element: Dict, styles: Dict):
        """Add a table row"""
        if not self.current_table or self.current_row_index >= len(
            self.current_table.rows
        ):
            return

        self.current_row = self.current_table.rows[self.current_row_index]
        self.current_cell_index = 0

        # Process row content
        self._convert_elements(element.get("children", []), styles)

        self.current_row_index += 1
        self.current_row = None

    def _add_table_cell(self, element: Dict, styles: Dict):
        """Add content to a table cell"""
        if not self.current_row or self.current_cell_index >= len(
            self.current_row.cells
        ):
            return

        cell = self.current_row.cells[self.current_cell_index]
        tag_name = element.get("tag", "").lower()
        attributes = element.get("attributes", {})

        # Handle colspan and rowspan
        colspan = int(attributes.get("colspan", 1))
        # rowspan = int(attributes.get("rowspan", 1))

        # Apply cell styles
        self._apply_cell_styles(cell, attributes, styles, tag_name == "th")

        # Save current paragraph context and switch to cell context
        saved_paragraph = self.current_paragraph
        self.current_paragraph = (
            cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        )

        # Process cell content
        self._convert_elements(element.get("children", []), styles)

        # Restore paragraph context
        self.current_paragraph = saved_paragraph
        self.current_cell_index += colspan

    def _apply_table_styles(self, table, attributes: Dict, styles: Dict):
        """Apply styles to the table"""
        # Apply auto-fit by default
        table.autofit = True

        # Apply specific width if provided
        width_attr = attributes.get("width") or attributes.get("style", "")
        if "width" in width_attr:
            width_match = re.search(r"width:\s*(\d+)(px|%)", width_attr)
            if width_match:
                width_value = int(width_match.group(1))
                if width_match.group(2) == "%":
                    # Convert percentage to approximate width (Word doesn't support % directly)
                    table_width = Inches(6)  # Approximate page width
                    width_value = int(table_width * width_value / 100)
                # table.width = WidthDocx(width_value)  # Would need proper width handling

        # Apply border styles
        if "border" in attributes.get("style", ""):
            self._apply_table_borders(table)

    def _apply_table_borders(self, table):
        """Apply borders to table"""
        try:
            tbl = table._tbl
            tblPr = tbl.tblPr

            # Add table borders
            tblBorders = OxmlElement("w:tblBorders")

            for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "auto")
                tblBorders.append(border)

            tblPr.append(tblBorders)
        except Exception as e:
            print(f"Table border styling failed: {e}")

    def _apply_cell_styles(
        self, cell, attributes: Dict, styles: Dict, is_header: bool = False
    ):
        """Apply styles to table cell"""
        # Set header styling
        if is_header:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Apply background color
        bg_color = self._extract_background_color(attributes, styles)
        if bg_color:
            try:
                shading_elm = parse_xml(f'<w:shd {qn("w:fill")}="{bg_color}"/>')
                cell._tc.get_or_add_tcPr().append(shading_elm)
            except Exception:
                pass

        # Apply text alignment
        align = attributes.get("align") or self._extract_text_align(attributes, styles)
        if align:
            for paragraph in cell.paragraphs:
                if align == "center":
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == "right":
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif align == "justify":
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _extract_background_color(self, attributes: Dict, styles: Dict) -> str:
        """Extract background color from attributes and styles"""
        # Check inline style
        style_attr = attributes.get("style", "")
        bg_match = re.search(r"background-color:\s*(#[0-9a-fA-F]+|\w+)", style_attr)
        if bg_match:
            return bg_match.group(1)

        # Check class styles
        class_attr = attributes.get("class", "")
        if class_attr:
            for class_name in class_attr.split():
                css_selector = f".{class_name}"
                if (
                    css_selector in styles
                    and "background-color" in styles[css_selector]
                ):
                    return styles[css_selector]["background-color"]

        return None

    def _extract_text_align(self, attributes: Dict, styles: Dict) -> str:
        """Extract text alignment from attributes and styles"""
        # Check inline style
        style_attr = attributes.get("style", "")
        align_match = re.search(r"text-align:\s*(\w+)", style_attr)
        if align_match:
            return align_match.group(1)

        # Check class styles
        class_attr = attributes.get("class", "")
        if class_attr:
            for class_name in class_attr.split():
                css_selector = f".{class_name}"
                if css_selector in styles and "text-align" in styles[css_selector]:
                    return styles[css_selector]["text-align"]

        return None

    def _cleanup_table(self):
        """Clean up empty table rows or cells"""
        if not self.current_table:
            return

        # Remove completely empty rows
        rows_to_remove = []
        for i, row in enumerate(self.current_table.rows):
            if all(cell.text.strip() == "" for cell in row.cells):
                rows_to_remove.append(i)

        # Remove rows in reverse order to avoid index issues
        for i in sorted(rows_to_remove, reverse=True):
            try:
                self.current_table._tbl.remove(self.current_table.rows[i]._tr)
            except Exception:
                pass

    def _handle_grid_container(self, element: Dict, styles: Dict):
        """Handle grid container (simulate with table)"""
        attributes = element.get("attributes", {})
        style_attr = attributes.get("style", "")

        # Check if this is a grid container
        is_grid = "display:grid" in style_attr or "display: grid" in style_attr
        is_flex = "display:flex" in style_attr or "display: flex" in style_attr

        if is_grid:
            self._handle_css_grid(element, styles)
        elif is_flex:
            self._handle_flex_container(element, styles)
        else:
            self._convert_elements(element.get("children", []), styles)

    def _handle_css_grid(self, element: Dict, styles: Dict):
        """Simulate CSS Grid with a table"""
        attributes = element.get("attributes", {})
        style_attr = attributes.get("style", "")

        # Extract grid template columns
        grid_cols = 1
        grid_template_match = re.search(
            r"grid-template-columns:\s*(repeat\((\d+),\s*1fr\)|[\w\s\(\)]+)", style_attr
        )
        if grid_template_match:
            if "repeat" in grid_template_match.group(1):
                grid_cols = int(grid_template_match.group(2))
            else:
                # Count columns by splitting
                cols = grid_template_match.group(1).split()
                grid_cols = len(cols)

        # Count rows needed
        children = [
            child
            for child in element.get("children", [])
            if child.get("type") == "element"
        ]
        grid_rows = (len(children) + grid_cols - 1) // grid_cols

        if grid_rows > 0 and grid_cols > 0:
            # Create table to simulate grid
            table = self.doc.add_table(rows=grid_rows, cols=grid_cols)
            table.autofit = True

            # Fill table with grid items
            child_index = 0
            for row in table.rows:
                for cell in row.cells:
                    if child_index < len(children):
                        child = children[child_index]
                        # Process child element in cell
                        saved_paragraph = self.current_paragraph
                        self.current_paragraph = cell.paragraphs[0]
                        self._convert_element(child, styles)
                        self.current_paragraph = saved_paragraph
                        child_index += 1

    def _handle_flex_container(self, element: Dict, styles: Dict):
        """Handle flex container (simulate with table row)"""
        attributes = element.get("attributes", {})

        # Create a single-row table to simulate flex container
        children = [
            child
            for child in element.get("children", [])
            if child.get("type") == "element"
        ]

        if children:
            table = self.doc.add_table(rows=1, cols=len(children))
            table.autofit = True

            for i, child in enumerate(children):
                cell = table.rows[0].cells[i]
                saved_paragraph = self.current_paragraph
                self.current_paragraph = cell.paragraphs[0]
                self._convert_element(child, styles)
                self.current_paragraph = saved_paragraph
