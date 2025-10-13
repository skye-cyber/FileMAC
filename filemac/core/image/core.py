import shutil
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import re
from pathlib import Path
from docx.shared import Inches, Mm
from docx import Document
import os
import sys
from tqdm import tqdm
from PIL import Image
import cv2
from typing import List, Tuple, Union, Optional
from ...utils.simple import logger
from ...utils.decorators import Decorators
from ...utils.formats import SUPPORTED_IMAGE_FORMATS
from ...utils.file_utils import modify_filename_if_exists, DirectoryScanner
from ...utils.colors import fg, rs

RESET = rs


class ImageCompressor:
    def __init__(self, input_image_path):
        self.input_image_path = input_image_path

    def resize_image(self, target_size):
        try:
            input_image_path = self.input_image_path
            ext = input_image_path[-3:]
            output_image_path = (
                os.path.splitext(input_image_path)[0] + f"_resized.{ext}"
            )

            original_image = Image.open(input_image_path)
            original_size = original_image.size
            size = os.path.getsize(input_image_path)
            print(f"Original image size {fg.YELLOW}{size / 1000_000:.2f}MiB{RESET}")

            # Calculate the aspect ratio of the original image
            aspect_ratio = original_size[0] / original_size[1]

            # Convert the target sixze to bytes
            tz = int(target_size[:-2])
            if target_size[-2:].lower() == "mb":
                target_size_bytes = tz * 1024 * 1024
            elif target_size[-2:].lower() == "kb":
                target_size_bytes = tz * 1024
            else:
                logger.warning(
                    f"Invalid units. Please use either {fg.BMAGENTA}'MB'{RESET}\
        or {fg.BMAGENTA}'KB'{RESET}"
                )

            # Calculate the new dimensions based on the target size
            new_width, new_height = ImageCompressor.calculate_new_dimensions(
                original_size, aspect_ratio, target_size_bytes
            )
            print(f"{fg.BLUE}Processing ..{RESET}")
            resized_image = original_image.resize((new_width, new_height))
            resized_image.save(output_image_path, optimize=True, format="png")
            t_size = os.path.getsize(output_image_path) / 1000_000
            print(f"{fg.BGREEN}Ok{RESET}")
            print(
                f"Image resized to {fg.BYELLOW}{t_size:.2f}{RESET} and saved to {fg.BYELLOW}{output_image_path}"
            )
        except KeyboardInterrupt:
            print("\nQuit⏹️")
            sys.exit(1)
        except KeyError:
            print("KeyError")
        except Exception as e:
            print(f"{fg.RED}{e}{RESET}")

    def calculate_new_dimensions(original_size, aspect_ratio, target_size_bytes):
        try:
            # Calculate the new dimensions based on the target size in bytes
            original_size_bytes = (
                original_size[0] * original_size[1] * 3
            )  # Assuming 24-bit color depth
            scale_factor = (target_size_bytes / original_size_bytes) ** 0.5

            new_width = int(original_size[0] * scale_factor)
            new_height = int(original_size[1] * scale_factor)

            return new_width, new_height
        except KeyboardInterrupt:
            print("\nQuit⏹️")
            sys.exit(1)
        except KeyError:
            print("KeyError")
        except Exception as e:
            print(f"{fg.RED}{e}{RESET}")


class ImageConverter:
    """Convert images file to from one format to another"""

    def __init__(self, input_file, out_format):
        self.input_file = input_file
        self.out_format = out_format

    def preprocess(self) -> list:
        try:
            files_to_process = []

            if os.path.isfile(self.input_file):
                files_to_process.append(self.input_file)
            elif os.path.isdir(self.input_file):
                if os.listdir(self.input_file) is None:
                    print("Cannot work with empty folder")
                    sys.exit(1)
                for file in os.listdir(self.input_file):
                    file_path = os.path.join(self.input_file, file)
                    if os.path.isfile(file_path):
                        files_to_process.append(file_path)

            return files_to_process
        except FileNotFoundError:
            print("File not found❕")
            sys.exit(1)

    def convert_image(self) -> os.PathLike:
        try:
            input_list = self.preprocess()
            out_f = self.out_format.upper()
            out_f = "JPEG" if out_f == "JPG" else out_f
            input_list = [
                item
                for item in input_list
                if any(
                    item.lower().endswith(ext)
                    for ext in SUPPORTED_IMAGE_FORMATS.values()
                )
            ]
            for file in tqdm(input_list):
                if out_f.upper() in SUPPORTED_IMAGE_FORMATS:
                    _ = os.path.splitext(file)[0]
                    output_filename = _ + SUPPORTED_IMAGE_FORMATS[out_f].lower()
                else:
                    print("Unsupported output format")
                    sys.exit(1)
                """Load the image using OpenCV: """
                img = cv2.imread(file)
                """Convert the OpenCV image to a PIL image: """
                pil_img = Image.fromarray(cv2.cvtColor(img, cv2.COLOR_BGR2RGB))

                pil_img.save(output_filename, out_f)

                print(f"Saved image as: {fg.DCYAN}{output_filename}{RESET}")

            return output_filename
        except KeyboardInterrupt:
            print("\nQuit❕")
            sys.exit(1)
        except AssertionError:
            print("Assertion failed.")
        except KeyError:
            print(
                f"{fg.RED}ERROR:\tPending Implementation for{fg.ICYAN} {out_f} {fg.BWHITE}format{RESET}"
            )
        except Exception as e:
            print(f"{fg.RED}{e}{RESET}")


class GrayscaleConverter:
    """
    Class for converting images to grayscale and saving the processed output.

    Attributes:
        input_obj (Optional[Union[list[str], str, os.PathLike]]): Input file(s) or directory.
        output_file (Optional[Union[list[str], str, os.PathLike]]): Output file path or directory.
    """

    def __init__(
        self,
        input_obj: Optional[Union[list[str], str, os.PathLike]],
        output_file: Optional[Union[list[str], str, os.PathLike]] = None,
    ):
        """
        Initializes the GrayscaleConverter object.

        Args:
            input_obj: Input file(s) or directory.
            output_file: Output file path or directory.
        """
        self.input_obj = input_obj
        self.output_file = output_file

    def get_output_file(
        self, image_path: Optional[Union[str, os.PathLike]] = None
    ) -> Union[str, os.PathLike]:
        """
        Computes the correct output file path for a given input file.

        Args:
            image_path: Path to the input file.

        Returns:
            The computed output file path.
        """
        logger.info(f"{fg.BWHITE}Obtaining output file name{RESET}")
        if self.output_file and self.output_file.endswith(
            tuple(SUPPORTED_IMAGE_FORMATS.values())
        ):
            return os.path.abspath(self.output_file)
        if self.output_file:
            return os.path.abspath(os.path.splitext(self.output_file)[0] + ".png")
        if image_path:
            return os.path.abspath(
                os.path.splitext(os.path.basename(image_path))[0] + ".png"
            )
        return "default_output.txt"

    def run(self):
        """
        Runs the image to grayscale conversion operation on the input files.

        Applies the for_loop_decorator to process each image in the input list.
        """
        file_list = DirectoryScanner(self.input_obj).run()

        @Decorators().for_loop_decorator(file_list)
        def process_image(self, image_path):
            """Processes a single image, converting it to grayscale and saving."""
            try:
                logger.info(f"{fg.YELLOW}Processing {fg.CYAN}{image_path}{RESET}")
                img = cv2.imread(image_path)
                if img is None:
                    raise FileNotFoundError(f"Could not read image: {image_path}")
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                _, thresh = cv2.threshold(
                    gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU
                )
                self.save_pil_image(thresh, image_path)
            except FileNotFoundError as e:
                logger.error(f"{fg.RED}{e}{RESET}")
            except Exception as e:
                raise
                logger.error(f"An unexpected error occurred: {fg.RED}{e}{RESET}")

        process_image(self)

    def save_pil_image(self, thresh, image_path):
        """
        Saves a NumPy array representing a grayscale image as a PIL Image.

        Args:
            thresh: The NumPy array representing the grayscale image.
            image_path: The path of the original image, used to derive the output filename.
        """
        try:
            img_pil = Image.fromarray(thresh)
            filename = self.get_output_file(image_path)
            filename = modify_filename_if_exists(filename)
            img_pil.save(filename)
            logger.info(f"{fg.GREEN}Image saved as {fg.BLUE}{filename}{RESET}")
        except Exception as e:
            raise
            logger.error(f"Unable to save the image: {fg.RED}{e}{RESET}")


class ImageDocxConverter:
    """
    A class for converting images to DOCX documents.
    """

    def __init__(
        self,
        image_list: list = None,
        input_dir: Union[str, os.PathLike] = None,
        output_path: Union[str, os.PathLike] = None,
        image_size: Tuple[float, float] = (6, 8),  # Default to 6x8 inches
        margin_mm: float = 25,  # Default margin of 25mm (approx 1 inch)
    ) -> None:
        """
        Initializes the ImageToDocxConverter object.

        Args:
            output_path: Path to save the output DOCX file + the file name e.g ~/Document/output.docx.
            filename: Name of the output DOCX file.
            image_size: Tuple (width, height) in inches.
            margin_mm: Margin in millimeters.
        """
        self.image_list = image_list
        self.input_dir = input_dir
        self.output_path = output_path if output_path else self.ensure_output_file()
        self.image_size = image_size
        self.margin_mm = margin_mm
        self.document = Document()  # Create a new document object filename

        # Set document margins in the constructor
        sections = self.document.sections
        for section in sections:
            section.top_margin = Mm(self.margin_mm)
            section.bottom_margin = Mm(self.margin_mm)
            section.left_margin = Mm(self.margin_mm)
            section.right_margin = Mm(self.margin_mm)
        self.create_output_directory()  # Create output directory in constructor

    def ensure_output_file(self) -> os.PathLike:
        file_name = "filemac_image2docx.docx"
        if self.input_dir:
            base_dir = self.input_dir
        else:
            base_dir = Path(self.image_list[0]).parent

        file_path = os.path.join(base_dir, file_name)

        return file_path

    def create_output_directory(self) -> None:
        """
        Creates the output directory if it does not exist.
        """
        Path(self.output_path).parent.mkdir(parents=True, exist_ok=True)

    def get_valid_images(self, image_paths: List[str]) -> List[str]:
        """
        Filters the list of image paths, returning only those with supported formats.

        Args:
            image_paths: A list of file paths to images.

        Returns:
            A list of file paths to valid images.
        """
        valid_images = []
        for image_path in image_paths:
            try:
                if Image.open(image_path).format.lower() in [
                    _formats[1:] for _formats in SUPPORTED_IMAGE_FORMATS.values()
                ]:
                    valid_images.append(image_path)
                else:
                    print(
                        f"{fg.MAGENTA}Skipping unsupported image format: {fg.CYAN}{image_path}{RESET}"
                    )
            except Exception as e:
                print(
                    f"{fg.RED}Error processing image {fg.YELLOW}{image_path} - {fg.RED} {e}{RESET}"
                )
        return valid_images

    def convert_images_to_docx(self, image_paths: List[str]) -> os.PathLike:
        """
        Converts a list of images to a single DOCX document.

        Args:
            image_paths: List of image file paths.
        """

        valid_images = self.get_valid_images(image_paths)
        if not valid_images:
            print("No valid images to convert.")
            return

        for image_path in valid_images:
            try:
                # Add a paragraph for each image
                paragraph = self.document.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(
                    image_path,
                    width=Inches(self.image_size[0]),
                    height=Inches(self.image_size[1]),
                )
                # Add a page break after each image, except the last one
                if image_path != valid_images[-1]:
                    self.document.add_page_break()
            except Exception as e:
                print(
                    f"{fg.RED}Error processing image {fg.YELLOW}{image_path}:{fg.RED} {e}{RESET}"
                )

        docx_file_path = (
            self.output_path
            if self.output_path.endswith(("docx", "doc"))
            else f"{self.output_path}.docx"
        )
        self.document.save(docx_file_path)
        return docx_file_path

    def convert_images_in_directory(self, input_dir, output_path) -> os.PathLike:
        """
        Converts all images in a directory to a PDF.

        Args:
            input_dir (str): The directory containing the images.
            output_path (str): The path to save the generated Word File.
            file_extensions (tuple, optional): Tuple of image file extensions to include.
        """

        if not os.path.exists(input_dir):
            raise FileNotFoundError(f"Directory not found: {input_dir}")

        image_paths = sorted(
            [os.path.join(input_dir, f) for f in os.listdir(input_dir)]
        )

        image_paths = self.get_valid_images(image_paths)

        if not image_paths:
            raise ValueError(f"No images found in directory: {input_dir}")

        self.create_pdf_from_images(image_paths, output_path)
        return output_path

    def run(self) -> os.PathLike:
        """
        Runs the conversion process.

        Args:
            image_paths: List of image file paths to convert.
        """
        if not any((self.image_list, self.input_dir)):
            print("No image paths provided.")
            sys.exit()

        if self.image_list and self.output_path:
            if all(os.path.exists(img) for img in self.image_list):
                docx_file_path = self.convert_images_to_docx(self.image_list)
        elif self.input_dir and self.output_path:
            if os.path.exists(self.input_dir):
                docx_file_path = self.convert_images_in_directory(
                    self.input_dir, self.output_path
                )

        if docx_file_path:
            print(
                f"{fg.GREEN_RG}Successfully created DOCX: {fg.BLUE}{docx_file_path}{RESET}"
            )

        return docx_file_path

    def cli(self, args: List[str]) -> None:
        """
        Main function to parse command line arguments and perform the conversion.

        Args:
            args: List of command line arguments.
        """
        if not args or "-h" in args or "--help" in args:
            print(
                """
                Usage: python image_to_docx.py [options] image1 image2 ... imageN

                Options:
                    -h, --help            show this help message and exit
                    -o, --output PATH     path to save the output DOCX file (default: current directory)
                    -n, --name FILENAME   name of the output DOCX file (default: output_document)
                    -s, --size WIDTHxHEIGHT  size of images in inches (e.g., 6x8) (default: 6x8)
                    -m, --margin MARGIN_MM margin in millimeters (default: 25)
                """
            )
            sys.exit()

        image_paths = []
        output_path = "."  # Current directory
        filename = "output_document"
        image_size = (6, 8)  # Default 6x8 inches
        margin_mm = 25

        i = 1
        while i < len(args):
            if args[i] in ("-o", "--output"):
                output_path = args[i + 1]
                i += 2
            elif args[i] in ("-n", "--name"):
                filename = args[i + 1]
                i += 2
            elif args[i] in ("-s", "--size"):
                try:
                    size_str = args[i + 1]
                    width, height = map(float, size_str.split("x"))
                    image_size = (width, height)
                except ValueError:
                    print("Invalid size format. Please use WIDTHxHEIGHT (e.g., 6x8).")
                    sys.exit(1)
                i += 2
            elif args[i] in ("-m", "--margin"):
                try:
                    margin_mm = float(args[i + 1])
                except ValueError:
                    print("Invalid margin format. Please provide a numeric value.")
                    sys.exit(1)
                i += 2
            else:
                if not args[i].startswith("-"):
                    image_paths.append(args[i])
                    i += 1
                else:
                    print(f"Unknown argument: {args[i]}")
                    sys.exit(1)

        converter = ImageDocxConverter(output_path, filename, image_size, margin_mm)
        converter.run(image_paths)


class ImagePdfConverter:
    """
    A class for converting images to PDF.
    """

    def __init__(
        self,
        image_list: list = None,
        input_dir=None,
        output_pdf_path=None,
        page_size=letter,
        order: bool = False,
        base: bool = False,
        walk: bool = False,
        clean: bool = False,
    ):
        self.image_list = image_list
        self.input_dir = input_dir
        self.page_size = page_size
        self.order = order
        self.base = base
        self.walk = walk
        self.clean = clean
        self.output_pdf_path = (
            output_pdf_path if output_pdf_path else self.ensure_output_file()
        )

    def ensure_output_file(self) -> os.PathLike:
        file_name = "filemac_image2pdf.pdf"
        if self.input_dir:
            base_dir = self.input_dir
            if self.base:
                one_file = os.listdir(self.input_dir)[0]
                base_name, ext = os.path.splitext(one_file)
                if "_img_" in base_name:
                    base_name = base_name.split("_img_")[0]
                    file_name = base_name + ".pdf"
                else:
                    file_name = self.input_dir.split("_imgs")[0] + ".pdf"
        else:
            base_dir = Path(self.image_list[0]).parent

        file_path = os.path.join(base_dir, file_name)

        return file_path

    def _clean(self, dirs: list):
        print(f"{fg.UWHITE}{fg.BWHITE}Clean Images Host dir{fg.RESET}")
        for d in dirs:
            abspath = os.path.abspath(d)
            print(f"{fg.BWHITE}Nuke: {fg.BYELLOW}{abspath}{fg.RESET}")
            # print(Path(d).is_relative_to(os.path.expanduser("~")))
            if (
                os.path.exists(d) and os.path.isdir(d)
                # and Path(d).is_relative_to(os.path.expanduser("~"))
            ):
                shutil.rmtree(abspath)

    def create_pdf_from_images(
        self, image_paths, output_pdf_path, resize_to_fit=True
    ) -> os.PathLike:
        """
        Creates a PDF from a list of image paths.

        Args:
            image_paths (list): A list of image file paths.
            output_pdf_path (str): The path to save the generated PDF.
            resize_to_fit (bool, optional): Whether to resize images to fit the page. Defaults to True.

        Raises:
            FileNotFoundError: If any image path is invalid.
            ValueError: If image_paths is empty or contains non-image files.
            Exception: for pillow image opening errors, or reportlab canvas errors.
        """

        if not image_paths:
            raise ValueError("Image paths list is empty.")

        for image_path in image_paths:
            if not os.path.exists(image_path):
                raise FileNotFoundError(f"Image not found: {image_path}")
            try:
                Image.open(image_path)
            except Exception as e:
                raise ValueError(f"Error opening image {image_path}: {e}")

        try:
            c = canvas.Canvas(output_pdf_path, pagesize=self.page_size)
            width, height = self.page_size

            for image_path in image_paths:
                img = Image.open(image_path)
                img_width, img_height = img.size

                if resize_to_fit:
                    ratio = min(width / img_width, height / img_height)
                    new_width = img_width * ratio
                    new_height = img_height * ratio
                    x = (width - new_width) / 2
                    y = (height - new_height) / 2
                else:
                    x = (width - img_width) / 2
                    y = (height - img_height) / 2
                    new_width = img_width
                    new_height = img_height

                c.drawImage(
                    image_path,
                    x,
                    y,
                    width=new_width,
                    height=new_height,
                    preserveAspectRatio=True,
                )
                c.showPage()

            c.save()

            return output_pdf_path
        except Exception as e:
            raise Exception(f"Error creating PDF: {e}")

    @staticmethod
    def ensure_format(input_image) -> os.PathLike:
        from ..imagepy.converter import ImageConverter

        converter = ImageConverter(input_image, "png")
        output_image = converter.convert_image()
        return output_image

    def extract_img_number(self, filename):
        match = re.search(r"_img_(\d+)", filename)
        return int(match.group(1)) if match else float("inf")

    def _sort(self, obj, ext):
        if self.order:
            if isinstance(obj, list):
                return sorted(
                    obj,
                    key=lambda f: self.extract_img_number(f),
                )
            return sorted(
                [
                    os.path.join(obj, f)
                    for f in os.listdir(obj)
                    if f.lower().endswith(ext)
                ],
                key=lambda f: self.extract_img_number(f),
            )
        else:
            return sorted(
                [
                    os.path.join(obj, f)
                    for f in os.listdir(obj)
                    if f.lower().endswith(ext)
                ]
            )

    def convert_images_in_directory_recursive(
        self, input_dir, output_pdf_path, file_extensions=(".jpg", ".jpeg", ".png")
    ):
        """
        Recursively walks through a directory and its subdirectories,
        converting images in each folder into a separate PDF.

        Args:
            input_dir (str): Root directory containing images.
            output_root (str): Directory to save the generated PDFs.
            file_extensions (tuple): Supported image extensions.
        """
        try:
            if not os.path.exists(input_dir):
                raise FileNotFoundError(f"Directory not found: {input_dir}")

            # if not os.path.exists(output_root):
            # os.makedirs(output_root)
            dclean = []
            for root, _, files in os.walk(input_dir):
                image_paths = [
                    os.path.join(root, f)
                    for f in files
                    if f.lower().endswith(file_extensions)
                ]

                if not image_paths:
                    continue  # No valid images in this directory

                # Optional: sort images with your custom logic
                image_paths = self._sort(image_paths, file_extensions)

                # Ensure formats are valid
                for index, image in enumerate(image_paths):
                    if not image.lower().endswith(file_extensions):
                        image_paths[index] = self.ensure_format(image)

                # Create a relative PDF name based on the subdir structure
                fname = os.path.split(root)[-1].split("_imgs")[0] + ".pdf"
                relative_path = os.path.join(
                    os.path.dirname((os.path.relpath(root, input_dir))), fname
                )
                # Host dir for images to be cleaned is clean is on
                dname = os.path.relpath(root, input_dir)
                dclean.append(dname)

                # pdf_name = relative_path.replace(os.sep, "_") + ".pdf"
                # pdf_output_path = os.path.join(output_root, pdf_name)

                # Create the PDF for this folder
                self.create_pdf_from_images(image_paths, relative_path)
                print(f"{fg.BWHITE}Created PDF{RESET}: {relative_path}")
            if self.clean:
                self._clean(dclean)
        except Exception as e:
            print(f"\033[31m{e}\033[0m")
            sys.exit(1)

    def convert_images_in_directory(
        self, input_dir, output_pdf_path, file_extensions=(".jpg", ".jpeg", ".png")
    ) -> os.PathLike:
        try:
            """
            Converts all images in a directory to a PDF.

            Args:
                input_dir (str): The directory containing the images.
                output_pdf_path (str): The path to save the generated PDF.
                file_extensions (tuple, optional): Tuple of image file extensions to include.
            """

            if not os.path.exists(input_dir):
                raise FileNotFoundError(f"Directory not found: {input_dir}")

            image_paths = self._sort(input_dir, ext=file_extensions)

            for index, image in enumerate(image_paths):
                if not image.endswith(file_extensions):
                    image_paths[index] = self.ensure_format(image)

            if not image_paths:
                raise ValueError(
                    f"\033[31mNo images found in directory:\033[1m {input_dir}\033[0m"
                )

            self.create_pdf_from_images(image_paths, output_pdf_path)
            return output_pdf_path
        except ValueError as e:
            print(e)
            sys.exit(1)

    def run(self) -> os.PathLike:
        """
        Runs the PDF creation based on the object's initialization parameters.
        """
        if self.image_list and self.output_pdf_path:
            if all(os.path.exists(img) for img in self.image_list):
                output_pdf_path = self.create_pdf_from_images(
                    self.image_list, self.output_pdf_path
                )
                print(f"{fg.GREEN}PDF created successfully from directory!{RESET}")
                print(
                    f"{fg.GREEN}Output:{RESET} {fg.BLUE}{output_pdf_path}{RESET}"
                )
            else:
                print(f"{fg.RED}One or more images in the list do not exist.{RESET}")
        elif self.input_dir and self.output_pdf_path:
            if os.path.exists(self.input_dir):
                if self.walk:
                    output_pdf_path = self.convert_images_in_directory_recursive(
                        self.input_dir, self.output_pdf_path
                    )
                else:
                    output_pdf_path = self.convert_images_in_directory(
                        self.input_dir, self.output_pdf_path
                    )
                    print(
                        f"{fg.GREEN}PDF created successfully from directory!{RESET}"
                    )
                    print(
                        f"{fg.BWHITE}Output:{RESET} {fg.BLUE}{output_pdf_path}{RESET}"
                    )
            else:
                print(
                    f"Directory {fg.YELLOW}{self.input_dir}{RESET} does not exist."
                )
        else:
            print(
                "Please provide either image_list and output_pdf_path or input_dir and output_pdf_path during object instantiation."
            )
            return
        return output_pdf_path
