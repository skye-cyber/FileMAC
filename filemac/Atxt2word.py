from docx import Document
from docx.shared import Pt, RGBColor


def text_to_word(text_file, output_file):
    # Create a new Document
    doc = Document()

    # Define formatting for headings and body text
    heading_font_name = 'Arial'
    heading_font_size = Pt(16)
    heading_font_color = RGBColor(0, 0, 255)  # Blue color

    body_font_name = 'Arial'
    body_font_size = Pt(12)
    body_font_color = RGBColor(0, 0, 0)  # Black color

    # Open the text file and read content
    with open(text_file, 'r') as file:
        lines = file.readlines()

    # Add paragraphs with different formatting
    for line in lines:
        # Check if the line is a heading (e.g., starts with a '#')
        if line.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])  # Remove the '# ' from the heading
            run.font.name = heading_font_name
            run.font.size = heading_font_size
            run.font.color.rgb = heading_font_color
            p.style = 'Heading1'
        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = body_font_name
            run.font.size = body_font_size
            run.font.color.rgb = body_font_color

    # Save the document
    doc.save(output_file)


if __name__ == "__main__":
    input_text_file = 'example.txt'
    output_word_file = 'output.docx'

    # Call the function
    text_to_word(input_text_file, output_word_file)
    print(f"Text file converted to Word document: {output_word_file}")


'''1. Enhanced Heading Detection:
Multi-Level Headings: Consider supporting multiple levels of headings (e.g., ## Heading 2, ### Heading 3). You can adjust the script to recognize different levels and apply corresponding styles.
2. Text Formatting Tags:
Bold and Italic Text: Allow text formatting such as bold and italic in the text file. You could use specific markers (e.g., *bold*, _italic_) and parse these markers to apply the formatting.
3. Custom Styles:
Custom Styles for Headings: Define and use custom styles for headings and body text instead of directly setting font properties. This makes it easier to update styles across the document later.
4. Error Handling:
File Existence and Read Errors: Implement error handling to manage cases where the input file might not exist or cannot be read.
5. Text File Parsing Improvements:
Rich Text Files: Consider supporting rich text formats or Markdown for more complex formatting in the source text files.
6. Additional Formatting Options:
Paragraph Spacing: Add options to control spacing before and after paragraphs.
Indentation: Allow indentation settings for different levels of text or headings.
'''


def text_to_word(text_file, output_file):
    # Create a new Document
    doc = Document()

    # Define formatting for headings and body text
    heading_styles = {
        # Heading 1
        1: {'font_size': Pt(16), 'font_color': RGBColor(0, 0, 255)},
        # Heading 2
        2: {'font_size': Pt(14), 'font_color': RGBColor(0, 0, 200)},
        # Heading 3
        3: {'font_size': Pt(12), 'font_color': RGBColor(0, 0, 150)},
    }

    body_font_name = 'Arial'
    body_font_size = Pt(12)
    body_font_color = RGBColor(0, 0, 0)  # Black color

    # Open the text file and read content
    with open(text_file, 'r') as file:
        lines = file.readlines()

    for line in lines:
        # Determine heading level or body text
        if line.startswith('# '):
            level = line.count('#')
            level = min(level, 3)  # Support up to 3 levels of headings
            style = heading_styles.get(level, heading_styles[1])
            p = doc.add_paragraph()
            # Remove '#' and extra space
            run = p.add_run(line[level+1:].strip())
            run.font.size = style['font_size']
            run.font.color.rgb = style['font_color']
            p.style = f'Heading{level}'
        else:
            p = doc.add_paragraph()
            run = p.add_run(line.strip())
            run.font.name = body_font_name
            run.font.size = body_font_size
            run.font.color.rgb = body_font_color

    # Save the document
    doc.save(output_file)


if __name__ == "__main__":
    input_text_file = 'example.txt'
    output_word_file = 'output.docx'

    # Call the function
    text_to_word(input_text_file, output_word_file)
    print(f"Text file converted to Word document: {output_word_file}")
